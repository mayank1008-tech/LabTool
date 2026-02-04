import os
import sys
import subprocess
import time
import smtplib
import json
import threading
import tkinter as tk
from tkinter import filedialog
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# --- COLOR CODES ---
GREEN = '\033[92m'
CYAN = '\033[96m'
YELLOW = '\033[93m'
RED = '\033[91m'
BOLD = '\033[1m'
RESET = '\033[0m'

# --- CHECK LIBRARIES ---
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print(f"{RED}{BOLD}![ERROR] 'python-docx' library is missing.{RESET}")
    sys.exit(1)

CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".lab_tool_config.json")
SUPPORTED_EXTS = ['.java', '.py', '.c', '.cpp', '.cs', '.js']

CURRENT_SETTINGS = {
    "h_font": "Times New Roman", "h_size": 16, "h_bold": True, "h_underline": False,
    "l_font": "Times New Roman", "l_size": 14, "l_bold": True, "l_underline": True,
    "b_font": "Times New Roman", "b_size": 12, "b_bold": False, "b_underline": False,
    "c_font": "Times New Roman", "c_size": 12, "c_bold": False, "c_underline": False
}

def print_banner():
    os.system('cls' if os.name == 'nt' else 'clear')
    banner = f"""
{GREEN}{BOLD}
██╗      █████╗ ██████╗ ████████╗ ██████╗  ██████╗ ██╗     
██║     ██╔══██╗██╔══██╗╚══██╔══╝██╔═══██╗██╔═══██╗██║     
██║     ███████║██████╔╝   ██║   ██║   ██║██║   ██║██║     
██║     ██╔══██║██╔══██╗   ██║   ██║   ██║██║   ██║██║     
███████╗██║  ██║██████╔╝   ██║   ╚██████╔╝╚██████╔╝███████╗
╚══════╝╚═╝  ╚═╝╚═════╝    ╚═╝    ╚═════╝  ╚═════╝ ╚══════╝
{RESET}
{CYAN}----------------------------------------------------------
{BOLD}Developer:{RESET} Mayank Jain
{BOLD}Github:   {RESET} https://github.com/mayank1008-tech
{BOLD}LinkedIn: {RESET} https://www.linkedin.com/in/mayank-jain-78a6bb321/
{BOLD}Insta:    {RESET} @mank_1008
{CYAN}----------------------------------------------------------{RESET}
    """
    print(banner)

def open_file_cross_platform(filename):
    if sys.platform == 'win32':
        os.startfile(filename)
    elif sys.platform == 'darwin':
        subprocess.call(['open', filename])
    else:
        subprocess.call(['xdg-open', filename])

def read_source_safe(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f: return f.read()
    except UnicodeDecodeError:
        try:
            with open(filepath, 'r', encoding='latin-1') as f: return f.read()
        except Exception as e: return f"// Error reading file: {str(e)}"

def send_mail_logic(file_path, receiver):
    try:
        with open(CONFIG_FILE, "r") as f: config = json.load(f)
        msg = MIMEMultipart()
        msg['From'] = config['sender_email']; msg['To'] = receiver
        msg['Subject'] = f"Lab Report: {os.path.basename(file_path)}"

        with open(file_path, "rb") as f:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(f.read())
        encoders.encode_base64(part); part.add_header('Content-Disposition', f"attachment; filename= {os.path.basename(file_path)}")
        msg.attach(part)

        server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
        server.login(config['sender_email'], config['app_key'])
        server.sendmail(config['sender_email'], receiver, msg.as_string()); server.quit()
        return True
    except Exception as e:
        return str(e)

def background_email_task(file_path, receiver):
    result = send_mail_logic(file_path, receiver)
    if result is True:
        print(f"\n{GREEN}{BOLD}[BG-EMAIL] Transmission to {receiver} successful! ✔{RESET}\n>> ", end="")
    else:
        print(f"\n{RED}[BG-EMAIL] Failed: {result}{RESET}\n>> ", end="")

def set_run_style(run, font, size, bold, underline):
    run.font.name = font; run.font.size = Pt(size); run.bold = bold; run.font.underline = underline

def generate_report():
    print_banner()
    source_file = input(f"\n{CYAN}{BOLD}>> Filename: {RESET}").strip()
    if not any(source_file.endswith(ext) for ext in SUPPORTED_EXTS) or not os.path.exists(source_file):
        print(f"{RED}![ERROR] Invalid file.{RESET}"); time.sleep(2); return

    aim = input(f"{CYAN}{BOLD}>> Enter Aim: {RESET}").strip()
    doc = Document()
    p = doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run_h = p.add_run("Program")
    set_run_style(run_h, CURRENT_SETTINGS['h_font'], CURRENT_SETTINGS['h_size'], CURRENT_SETTINGS['h_bold'], CURRENT_SETTINGS['h_underline'])

    for label, content_type in [("Aim:- ", "aim"), ("SOURCE CODE:-", "code"), ("OUTPUT:-", "out")]:
        p = doc.add_paragraph()
        l_run = p.add_run(label)
        set_run_style(l_run, CURRENT_SETTINGS['l_font'], CURRENT_SETTINGS['l_size'], CURRENT_SETTINGS['l_bold'], CURRENT_SETTINGS['l_underline'])

        if content_type == "aim":
            b_run = p.add_run(aim)
            set_run_style(b_run, CURRENT_SETTINGS['b_font'], CURRENT_SETTINGS['b_size'], CURRENT_SETTINGS['b_bold'], CURRENT_SETTINGS['b_underline'])
        elif content_type == "code":
            code = read_source_safe(source_file)
            c_run = doc.add_paragraph().add_run(code)
            set_run_style(c_run, CURRENT_SETTINGS['c_font'], CURRENT_SETTINGS['c_size'], CURRENT_SETTINGS['c_bold'], CURRENT_SETTINGS['c_underline'])
        elif content_type == "out":
            p_out = doc.add_paragraph(); p_out.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p_out.add_run("\n\n[ PASTE SCREENSHOT HERE ]\n\n"); r.font.color.rgb = RGBColor(180, 180, 180); r.italic = True

    out_name = os.path.splitext(source_file)[0] + "_Report.docx"

    while True:
        try:
            doc.save(out_name)
            print(f"{GREEN}✔ Document Generated Locally.{RESET}")

            # --- NEW: REMINDER BEFORE OPENING ---
            print(f"{YELLOW}[REMINDER] Please paste your OUTPUT SCREENSHOTS in the Word file now!{RESET}")
            time.sleep(1.5)
            # ------------------------------------

            open_file_cross_platform(out_name)

            if input(f"\n{BOLD}>> Email file? (y/n): {RESET}").lower() == 'y':
                print(f"{YELLOW}>> Opening File Selector...{RESET}")

                root = tk.Tk(); root.withdraw(); root.attributes('-topmost', True)
                selected_path = filedialog.askopenfilename(
                    title="Select Report to Email",
                    filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
                )

                if not selected_path:
                    print(f"{RED}![CANCEL] No file selected.{RESET}")
                    break

                if not os.path.exists(CONFIG_FILE):
                    key = input("Enter App Key: ")
                    with open(CONFIG_FILE, "w") as f: json.dump({"sender_email": "cr.lab.bot.word@gmail.com", "app_key": key}, f)

                receiver = input(f"{CYAN}>> Destination Email: {RESET}").strip()
                bg_choice = input(f"{BOLD}>> Send in background? (y/n): {RESET}").lower()

                print(f"{YELLOW}[!] Ensure file is CLOSED in Word.{RESET}")
                while True:
                    try:
                        with open(selected_path, "rb"): pass
                        break
                    except PermissionError:
                        print(f"{RED}![ERROR] File is Locked by Word/OneDrive.{RESET}")
                        if input(f"{YELLOW}>> Close it and press ENTER to retry (or 'q' to quit): {RESET}").lower() == 'q':
                            return

                if bg_choice == 'y':
                    threading.Thread(target=background_email_task, args=(os.path.abspath(selected_path), receiver), daemon=True).start()
                    print(f"{GREEN}✔ Bot dispatched.{RESET}")
                else:
                    print(f"{YELLOW}[SMTP] Transmitting...{RESET}")
                    res = send_mail_logic(os.path.abspath(selected_path), receiver)
                    if res is True: print(f"{GREEN}✔ Transmission Successful!{RESET}")
                    else: print(f"{RED}![FAIL] {res}{RESET}")
            break
        except PermissionError:
            print(f"\n{RED}![ERROR] Access Denied. Close Word and try again.{RESET}"); input()
        except Exception as e:
            print(f"{RED}![ERROR] {e}{RESET}"); break

    input(f"\n{YELLOW}Press Enter to return to menu...{RESET}")

def main_menu():
    while True:
        print_banner()
        print(f"{GREEN}{BOLD}--- MAIN COMMAND CENTER ---{RESET}")
        print(f"{CYAN}[1]{RESET} Generate Report  {CYAN}[2]{RESET} Styles  {CYAN}[3]{RESET} Reset Bot  {CYAN}[4]{RESET} Exit")
        choice = input(f"\n{BOLD}Selection: {RESET}")
        if choice == '1': generate_report()
        elif choice == '2': customization_menu()
        elif choice == '3':
            if os.path.exists(CONFIG_FILE): os.remove(CONFIG_FILE)
            print(f"{GREEN}✔ Credentials Purged.{RESET}"); time.sleep(1)

        elif choice == '4':
            # --- THE CINEMATIC EXIT ---
            print()
            quote = "\"Trust the process & U Are the process\" - Mayank Jain"
            sys.stdout.write(f"{CYAN}{BOLD}")
            for char in quote:
                sys.stdout.write(char)
                sys.stdout.flush()
                time.sleep(0.04)
            sys.stdout.write(f"{RESET}\n")

            time.sleep(0.5)
            print(f"\n{RED}[SYSTEM] Shutting down", end="")

            # Dot animation (...)
            for _ in range(3):
                time.sleep(0.5)
                sys.stdout.write(".")
                sys.stdout.flush()

            print(f"{RESET}")
            time.sleep(0.5)
            break

def customization_menu():
    while True:
        print_banner()
        keys = list(CURRENT_SETTINGS.keys())
        for i, k in enumerate(keys, 1): print(f"{CYAN}[{i}]{RESET} {k}: {CURRENT_SETTINGS[k]}")
        choice = input(f"\n{BOLD}Selection (0 to back): {RESET}")
        if choice == '0': break
        try:
            k = keys[int(choice)-1]; val = input(f"New value for {k}: ")
            if 'size' in k: CURRENT_SETTINGS[k] = int(val)
            elif 'bold' in k or 'underline' in k: CURRENT_SETTINGS[k] = val.lower() == 'y'
            else: CURRENT_SETTINGS[k] = val
        except: pass

if __name__ == "__main__":
    main_menu()
